import uuid
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.colors import Color
import re

def json_to_markdown(resume_data: dict) -> str:
    md_content = []
    
    personal_info = resume_data.get('personal_info', {})
    
    if personal_info.get('name'):
        md_content.append(f"# {personal_info['name'].upper()}")
        md_content.append("")
        
    #contact
    contact = []
    if personal_info.get('email'):
        contact.append(f"Email: {personal_info['email']}")
    if personal_info.get('phone'):
        contact.append(f"Phone: {personal_info['phone']}")
    if personal_info.get('location'):
        contact.append(f"Location: {personal_info['location']}")
    if personal_info.get('linkedin'):
        contact.append(f"Linkedin: {personal_info['linkedin']}")
    if personal_info.get('github'):
        contact.append(f"Github: {personal_info['github']}")
    if personal_info.get('website'):
        contact.append(f"Website: {personal_info['website']}")
        
    if contact:
        md_content.append(" | ".join(contact))
        md_content.append("")
        md_content.append("---")
        md_content.append("")
    
    #summary
    if resume_data.get('summary'):
        md_content.append("### SUMMARY")
        md_content.append("---")
        md_content.append("")
        md_content.append(resume_data['summary'])
    
    #education
    if resume_data.get('education'):
        md_content.append("### Education")
        md_content.append("---")
        for edu in resume_data['education']:
            if not edu:
                continue
            
            parts = []
            if edu.get('degree'):
                parts.append(f"**{edu['degree']}**")
            if edu.get('institution'):
                parts.append(edu['institution'])
            if edu.get('year'):
                parts.append(str(edu['year']))
            
            if parts:
                md_content.append(" | ".join(parts))
                md_content.append("")
    
    #experience
    if resume_data.get('experience'):
        md_content.append('### Experience')
        md_content.append("---")
        md_content.append("")
        
        for exp in resume_data['experience']:
            if not exp:
                continue
            
            title = []
            if exp.get('title'):
                title.append(f"**{exp['title']}**")
            if exp.get('company'):
                title.append(exp['company'])
            if exp.get('duration'):
                title.append(exp['duration'])
            
            if title:
                md_content.append(" | ".join(title))
                md_content.append("")
            
            #job stuff
            if exp.get('description'):
                if isinstance(exp['description'], list):
                    for item in exp['description']:
                        md_content.append(f"• {item}")
                else:
                    md_content.append(f"• {exp['description']}")
                md_content.append("")
    
    #projects
    if resume_data.get('projects'):
        md_content.append("### Projects")
        md_content.append("---")
        md_content.append("")
        for project in resume_data['projects']:
            if not project:
                continue
                
            proj_lines = []
            if project.get('name'):
                proj_lines.append(f"**{project['name']}**")
            
            if project.get('technologies'):
                if isinstance(project['technologies'], list):
                    tech_list = [tech for tech in project['technologies']]
                    if tech_list:
                        proj_lines.append(f"Technologies: {', '.join(tech_list)}")
                else:
                    proj_lines.append(f"Technologies: {project['technologies']}")
            
            if project.get('url'):
                proj_lines.append(project['url'])
                
            if proj_lines:
                md_content.append(" | ".join(proj_lines))
                md_content.append("")
                
            if project.get('description'):
                md_content.append(f"• {project['description']}")
                md_content.append("")
                    
    #skills
    if resume_data.get('skills'):
        md_content.append("## Technical Skills")
        md_content.append("---")
        md_content.append("")
        
        skills = resume_data['skills']
        if isinstance(skills, list):
            categories = {
                'Languages': ['python', 'java', 'javascript', 'c++', 'c#', 'go', 'rust', 'typescript', 'swift', 'kotlin', 'sql', 'html', 'css'],
                'Frameworks/Libraries': ['react', 'angular', 'vue', 'django', 'flask', 'spring', 'express', 'next', 'node'],
                'Tools/Technologies': ['git', 'docker', 'kubernetes', 'aws', 'azure', 'jenkins', 'linux', 'mongodb', 'postgresql', 'mysql']
            }
            
            categorized = {cat: [] for cat in categories}
            categorized['Other'] = []
            
            for skill in skills:
                categorized_skill = False
                for category, keyword in categories.items():
                    if any(key in skill.lower() for key in keyword):
                        categorized[category].append(skill)
                        categorized_skill = True
                        break
                if not categorized_skill:
                    categorized['Other'].append(skill)
            
            for category, items in categorized.items():
                if items:
                    md_content.append(f"**{category}:** {', '.join(items)}")
                    md_content.append("")
            
            md_content.append("")
    
    #achievements
    if resume_data.get('achievements'):
        md_content.append("## Achievements")
        md_content.append("---")
        md_content.append("")
        
        for achievement in resume_data['achievements']:
            md_content.append(f"• {achievement}")
            
        md_content.append("")
    
    return "\n".join(md_content)


def markdown_to_docx(markdown_content: str) -> tuple[BytesIO, str]:

    doc = Document()
    lines = markdown_content.split("\n")
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if line.startswith('# '):
            name = line[2:]
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(name)
            name_run.bold = True
            name_run.font.size = Pt(18)
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif ('Email:' in line or 'Phone:' in line) and '|' in line:
            contact_para = doc.add_paragraph(line)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # handle horizontal rules (---)
        elif line.startswith('---'):
            hr_para = doc.add_paragraph('_' * 105)
            hr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # handle headings (###, ##)
        elif line.startswith('### '):
            heading_text = line[4:]
            doc.add_heading(heading_text, level=2)
            
        elif line.startswith('## '):
            heading_text = line[3:] 
            doc.add_heading(heading_text, level=2)
            
        # handle bullet points
        elif line.startswith('• '):
            bullet_text = line[2:]
            doc.add_paragraph(bullet_text, style='List Bullet')
        
        # Handle bold text with ** **
        elif '**' in line:
            para = doc.add_paragraph()
            parts = line.split('**')
            for j, part in enumerate(parts):
                if j % 2 == 0:
                    if part:
                        para.add_run(part)
                else:
                    if part:
                        run = para.add_run(part)
                        run.bold = True
        
        else:
            if line:
                doc.add_paragraph(line)
        
        i += 1
        
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    uid = str(uuid.uuid4())[:8]
    filename = f"resume_{timestamp}_{uid}.docx"
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer, filename

def markdown_to_pdf(markdown_content: str) -> tuple[BytesIO, str]:
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    content = []
    
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Title'],
        alignment=1, 
        fontSize=18,
        spaceAfter=12,
        fontName='Helvetica-Bold'
    )
    
    contact_style = ParagraphStyle(
        'ContactStyle', 
        parent=styles['Normal'],
        alignment=1,  
        spaceAfter=6,
        fontSize=10
    )
    
    docx_color = Color(0.31, 0.506, 0.741)
    heading_style = ParagraphStyle(
        'HeadingStyle',
        parent=styles['Heading2'],
        textColor=docx_color,
        fontSize=14,
        spaceAfter=2,  
        spaceBefore=12,
        fontName='Helvetica-Bold'
    )
    
    hr_style = ParagraphStyle(
        'HorizontalRule',
        parent=styles['Normal'],
        alignment=1,
        spaceAfter=4,
        spaceBefore=0  
    )
    
    project_title_style = ParagraphStyle(
        'ProjectTitleStyle',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=6,
        spaceBefore=13,
        fontName='Helvetica-Bold'
    )
    
    skills_style = ParagraphStyle(
        'SkillStyle',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=4,
        spaceBefore=2
    )
    
    lines = markdown_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Handle name heading
        if line.startswith('# '):
            name = line[2:]
            content.append(Paragraph(name, title_style))
            content.append(Spacer(1, 0.1*inch))
            
        # Handle contact info
        elif ('Email:' in line or 'Phone:' in line) and '|' in line:
            content.append(Paragraph(line, contact_style))
            
        # Handle horizontal rules (---)
        elif line.startswith('---'):
            content.append(Paragraph('_' * 82, hr_style))
            
        # Handle headings (###, ##)
        elif line.startswith('### '):
            heading_text = line[4:]
            content.append(Paragraph(heading_text, heading_style))
            
        elif line.startswith('## '):
            heading_text = line[3:] 
            content.append(Paragraph(heading_text, heading_style))
            
        # Handle bullet points
        elif line.startswith('• '):
            bullet_text = line[2:]
            content.append(Paragraph(f"• {bullet_text}", styles['Normal']))
            content.append(Spacer(1, 0.02*inch))
        
        # Handle bold text with ** ** 
        elif '**' in line:
            # Convert **text** to <b>text</b> for reportlab
            formatted_line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
            
            if ':' in line and any(skill in line for skill in ['Languages:', 'Frameworks:', 'Tools:', 'Other:']):
                skills_formatted = re.sub(r'\*\*(.*?):\*\*', r'<b>\1:</b>', line)
                content.append(Paragraph(skills_formatted, skills_style))
            else:
                content.append(Paragraph(formatted_line, project_title_style))
            
        else:
            if line:
                content.append(Paragraph(line, styles['Normal']))
                content.append(Spacer(1, 0.02*inch))
        
        i += 1
    
    content.append(Spacer(1, 0.1*inch))
    
    doc.build(content)
    buffer.seek(0)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    uid = str(uuid.uuid4())[:8]
    filename = f"resume_{timestamp}_{uid}.pdf"
    return buffer, filename
